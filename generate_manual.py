"""
Khalisni Platform – Comprehensive User Manual Generator  v2.0
Produces: D:\ghassan\Khalisni_User_Manual.docx
Covers every page and every action for all five roles.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG_DIR = r"D:\ghassan\user image"
OUT_PATH = r"D:\ghassan\Khalisni_User_Manual.docx"

BRAND_BLUE   = RGBColor(0x1A, 0x56, 0xDB)
DARK_NAVY    = RGBColor(0x1E, 0x29, 0x3B)
MID_GREY     = RGBColor(0x6B, 0x72, 0x80)
BODY_TEXT    = RGBColor(0x37, 0x41, 0x51)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_GREEN = RGBColor(0x05, 0x96, 0x69)
STEP_BLUE    = RGBColor(0x1D, 0x4E, 0xD8)


def img(name):
    return os.path.join(IMG_DIR, name)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_hr(doc, color_hex="1A56DB"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.color.rgb = DARK_NAVY
    run.font.size = Pt(20)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    add_hr(doc)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = BRAND_BLUE
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.color.rgb = DARK_NAVY
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)


def h4(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = STEP_BLUE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.color.rgb = BODY_TEXT
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_TEXT
    p.paragraph_format.left_indent  = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after  = Pt(3)


def step(doc, number, text):
    """Numbered step with bold step number."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"Step {number}:  ")
    r1.font.size = Pt(10.5)
    r1.font.bold = True
    r1.font.color.rgb = BRAND_BLUE
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BODY_TEXT
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(4)


def note(doc, text, kind="NOTE"):
    colors = {
        "NOTE":    ("EFF6FF", BRAND_BLUE),
        "WARNING": ("FEF9C3", RGBColor(0x92, 0x40, 0x0E)),
        "TIP":     ("F0FDF4", ACCENT_GREEN),
    }
    bg, fc = colors.get(kind, colors["NOTE"])
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f"  {kind}:  {text}")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = fc
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_image(doc, filename, caption="", width=Inches(5.8)):
    path = img(filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run().add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Figure: {caption}")
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = MID_GREY
        cap.paragraph_format.space_after = Pt(10)


def table(doc, rows_data, headers=None):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light List Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if headers:
        row = tbl.add_row()
        for i, h in enumerate(headers):
            c = row.cells[i]
            set_cell_bg(c, "1A56DB")
            r = c.paragraphs[0].add_run(h)
            r.font.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(10)
    for idx, (col1, col2) in enumerate(rows_data):
        row = tbl.add_row()
        if idx % 2 == 0:
            set_cell_bg(row.cells[0], "F8FAFC")
            set_cell_bg(row.cells[1], "F8FAFC")
        for cell, txt in zip(row.cells, [col1, col2]):
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(10)
            r.font.color.rgb = BODY_TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def status_table(doc, statuses):
    """Three-column status legend."""
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Light List Accent 1"
    hrow = tbl.add_row()
    for i, h in enumerate(["Status", "Display Name", "Meaning"]):
        c = hrow.cells[i]
        set_cell_bg(c, "1A56DB")
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for idx, (code, display, meaning) in enumerate(statuses):
        row = tbl.add_row()
        if idx % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, "F8FAFC")
        for cell, txt in zip(row.cells, [code, display, meaning]):
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(10)
            r.font.color.rgb = BODY_TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def pb(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════

def build_cover(doc):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "1A56DB")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("Khalisni  —  خلصني")
    r.font.size = Pt(38); r.font.bold = True; r.font.color.rgb = WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Comprehensive User Manual")
    r2.font.size = Pt(20); r2.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFF)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(50)
    r3 = p3.add_run("Service Request Management Platform")
    r3.font.size = Pt(14); r3.font.color.rgb = WHITE

    doc.add_paragraph()
    for k, v in [
        ("Version",      "2.0"),
        ("Date",         "June 2026"),
        ("Audience",     "Customers · Employees · Providers · Administrators"),
        ("Platform URL", "khalisni.raedaltabanjeh.com"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = p.add_run(f"{k}:  ")
        rb.font.bold = True; rb.font.color.rgb = BRAND_BLUE; rb.font.size = Pt(11)
        rv = p.add_run(v)
        rv.font.color.rgb = DARK_NAVY; rv.font.size = Pt(11)
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# TOC
# ═══════════════════════════════════════════════════════════════════════════

def build_toc(doc):
    h1(doc, "Table of Contents")
    entries = [
        ("1",      "Introduction & Platform Overview"),
        ("2",      "System Roles Overview"),
        ("3",      "Getting Started — Registration, Login & Account Setup"),
        ("  3.1",  "Creating a New Account (Registration)"),
        ("  3.2",  "Logging In"),
        ("  3.3",  "Forgot Password / Password Reset"),
        ("  3.4",  "Switching the Interface Language"),
        ("  3.5",  "Logging Out"),
        ("4",      "CUSTOMER — Complete Guide"),
        ("  4.1",  "Customer Dashboard"),
        ("  4.2",  "Browsing the Service Catalog"),
        ("  4.3",  "Searching & Filtering Services"),
        ("  4.4",  "Viewing a Service Detail Page"),
        ("  4.5",  "Submitting a New Service Request (Order)"),
        ("  4.6",  "Uploading Documents to an Order"),
        ("  4.7",  "Viewing My Orders List"),
        ("  4.8",  "Viewing an Order's Details & Timeline"),
        ("  4.9",  "Tracking an Order Without Logging In (Public Tracking)"),
        ("  4.10", "Responding to a Missing Documents Request"),
        ("  4.11", "Downloading a Final Document"),
        ("  4.12", "Managing Your Profile"),
        ("  4.13", "Viewing Notifications"),
        ("5",      "EMPLOYEE — Complete Guide"),
        ("  5.1",  "Employee Dashboard"),
        ("  5.2",  "Review Queue — Viewing & Filtering Orders"),
        ("  5.3",  "Opening an Order for Review"),
        ("  5.4",  "Starting a Review"),
        ("  5.5",  "Verifying & Approving Documents"),
        ("  5.6",  "Rejecting a Document"),
        ("  5.7",  "Requesting Missing Documents from the Customer"),
        ("  5.8",  "Adding an Internal Note"),
        ("  5.9",  "Assigning a Provider"),
        ("  5.10", "Sending a Manual Notification"),
        ("  5.11", "Returning an Order to the Provider"),
        ("  5.12", "Returning an Order to Internal Review"),
        ("  5.13", "Completing (Approving) an Order"),
        ("  5.14", "Rejecting an Order"),
        ("  5.15", "Employee Reports"),
        ("  5.16", "Missing Service Requests"),
        ("6",      "PROVIDER — Complete Guide"),
        ("  6.1",  "Provider Dashboard"),
        ("  6.2",  "Viewing Assigned Orders"),
        ("  6.3",  "Opening an Assigned Order"),
        ("  6.4",  "Updating Order Status"),
        ("  6.5",  "Uploading the Final Deliverable Document"),
        ("7",      "ADMINISTRATOR — Complete Guide"),
        ("  7.1",  "Admin Dashboard Overview"),
        ("  7.2",  "Orders Management"),
        ("  7.3",  "Rule Management Hub"),
        ("    7.3.1",  "Services & Pricing — Viewing the List"),
        ("    7.3.2",  "Services & Pricing — Creating a New Service"),
        ("    7.3.3",  "Services & Pricing — Editing a Service"),
        ("    7.3.4",  "Services & Pricing — Disabling a Service"),
        ("    7.3.5",  "Required Documents — Creating a Document Rule"),
        ("    7.3.6",  "Provider Assignment Rules"),
        ("    7.3.7",  "Workflow Rules"),
        ("    7.3.8",  "Notification Templates — Creating a Template"),
        ("    7.3.9",  "Payments Management"),
        ("    7.3.10", "System Settings"),
        ("    7.3.11", "Audit Logs"),
        ("    7.3.12", "Users & Roles — Creating a User"),
        ("    7.3.13", "Users & Roles — Editing a User's Permissions"),
        ("  7.4",  "Service Categories"),
        ("  7.5",  "Service Relations"),
        ("  7.6",  "Providers Management — Creating a Provider"),
        ("  7.7",  "Provider Service Assignments"),
        ("  7.8",  "Missing Service Requests (Admin)"),
        ("  7.9",  "Public Site Management"),
        ("    7.9.1",  "Homepage Content Editor"),
        ("    7.9.2",  "Advertisement Manager"),
        ("    7.9.3",  "Theme Settings"),
        ("    7.9.4",  "Preview Public Page"),
        ("  7.10", "Help Guide Management"),
        ("  7.11", "Admin Reports"),
        ("  7.12", "Notifications Management"),
        ("8",      "Order Status Reference"),
        ("9",      "Frequently Asked Questions"),
        ("10",     "Glossary"),
    ]
    for num, title in entries:
        p = doc.add_paragraph()
        indent = (len(num) - len(num.lstrip())) * 0.18
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_after = Pt(2)
        rn = p.add_run(num.strip().ljust(9))
        rn.font.size = Pt(10.5); rn.font.bold = True; rn.font.color.rgb = BRAND_BLUE
        rt = p.add_run(title)
        rt.font.size = Pt(10.5); rt.font.color.rgb = DARK_NAVY
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════

def ch1(doc):
    h1(doc, "1.  Introduction & Platform Overview")
    body(doc,
         "Khalisni (خلصني — meaning 'Get it done for me') is a Jordanian digital government "
         "service-request platform. Citizens submit applications for official services — "
         "passport renewals, civil status certificates, property registration deeds, business "
         "licences, and more — from any device. Staff review documents, providers execute "
         "the services, and administrators configure everything without writing a single line of code.")
    body(doc,
         "This manual is the complete operational reference for every role. Every screen, every "
         "button, and every step is documented with screenshots taken directly from the live "
         "staging environment.")
    h2(doc, "What You Can Do With Khalisni")
    bullet(doc, "Submit government and administrative service requests online — no queues, no paper.")
    bullet(doc, "Upload required documents securely and track their verification status in real time.")
    bullet(doc, "Follow your order through every stage with a timestamped timeline.")
    bullet(doc, "Track any order publicly using only an Order Number and phone number.")
    bullet(doc, "Receive in-app notifications at every key milestone.")
    bullet(doc, "Manage the entire platform — services, users, providers, content — from one admin panel.")
    h2(doc, "Technical Requirements")
    bullet(doc, "A modern web browser (Chrome, Edge, Firefox, Safari — latest two versions).")
    bullet(doc, "An active internet connection.")
    bullet(doc, "No software installation required.")
    note(doc, "All screenshots in this manual were captured from the staging environment at "
              "khalisnidev.raedaltabanjeh.com. Your production URL may differ.", "NOTE")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 2 — ROLES
# ═══════════════════════════════════════════════════════════════════════════

def ch2(doc):
    h1(doc, "2.  System Roles Overview")
    body(doc, "Khalisni uses a five-role access model. After login, each role sees a "
              "completely different dashboard and navigation menu tailored to their responsibilities.")
    table(doc, [
        ("Customer",  "Citizens who browse the service catalog, submit requests, upload documents, and track orders."),
        ("Employee",  "Internal government staff who review submitted orders, verify documents, and assign providers."),
        ("Support",   "Customer-facing agents with read access to orders plus access to service categories and relations."),
        ("Provider",  "External or internal service executors who receive assigned orders and upload final results."),
        ("Admin",     "Super-users who configure all platform settings: services, rules, users, content, and more."),
    ], headers=["Role", "Responsibilities"])
    note(doc, "A user can only have one role. Admins assign roles when creating accounts. "
              "The role cannot be changed by the user themselves.", "NOTE")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 3 — GETTING STARTED
# ═══════════════════════════════════════════════════════════════════════════

def ch3(doc):
    h1(doc, "3.  Getting Started")

    h2(doc, "3.1  Creating a New Account (Registration)")
    body(doc, "Any visitor can self-register as a Customer. Employee, Provider, Support, and "
              "Admin accounts must be created by an administrator.")
    step(doc, 1, "Open the platform URL in your browser.")
    step(doc, 2, "Click Register in the top navigation bar of the public home page.")
    step(doc, 3, "Fill in the registration form:")
    bullet(doc, "Full Name — your legal or display name.", 1)
    bullet(doc, "Email — used for login and notifications.", 1)
    bullet(doc, "Phone — primary contact number.", 1)
    bullet(doc, "Password — choose a strong password (at least 8 characters).", 1)
    step(doc, 4, "Click the Register / Create Account button.")
    step(doc, 5, "You are now logged in automatically as a Customer.")
    note(doc, "Passwords are stored using secure one-way hashing. Never share your password.", "WARNING")

    h2(doc, "3.2  Logging In")
    step(doc, 1, "Click Login in the top navigation bar.")
    step(doc, 2, "Enter your Email address and Password.")
    step(doc, 3, "Click Sign In.")
    step(doc, 4, "The system redirects you to your role-specific dashboard automatically.")
    note(doc, "Your session is stored in the browser tab only (sessionStorage). "
              "Closing the tab will log you out. Opening a new tab requires logging in again.", "NOTE")

    h2(doc, "3.3  Forgot Password / Password Reset")
    step(doc, 1, "On the Login page, click Forgot Password.")
    step(doc, 2, "Enter the email address associated with your account and click Send Reset Link.")
    step(doc, 3, "Check your inbox for the password reset email.")
    step(doc, 4, "Click the link in the email. You have a limited time window to use it.")
    step(doc, 5, "Enter and confirm your new password on the reset page.")
    step(doc, 6, "Click Save New Password.")
    step(doc, 7, "Return to the Login page and sign in with your new password.")

    h2(doc, "3.4  Switching the Interface Language")
    body(doc, "Khalisni is fully bilingual (Arabic and English). The interface direction "
              "switches automatically between RTL (Arabic) and LTR (English).")
    step(doc, 1, "Locate the language switcher in the top navigation bar (shows 'EN' or 'AR').")
    step(doc, 2, "Click it to toggle between Arabic and English.")
    step(doc, 3, "The entire interface — labels, buttons, and content — updates instantly.")
    note(doc, "Your language preference is saved for the current session. It resets on the next login.", "TIP")

    h2(doc, "3.5  Logging Out")
    step(doc, 1, "Click your name or avatar in the top-right corner of the dashboard.")
    step(doc, 2, "Click Log Out from the dropdown menu.")
    step(doc, 3, "You are returned to the public home page and your session is cleared.")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 4 — CUSTOMER
# ═══════════════════════════════════════════════════════════════════════════

def ch4(doc):
    h1(doc, "4.  Customer Guide")
    body(doc, "This chapter covers every action available to a logged-in Customer, from "
              "browsing services to downloading completed order results.")

    # 4.1 Dashboard
    h2(doc, "4.1  Customer Dashboard")
    body(doc, "After login, customers land on the Customer Dashboard (URL: /customer). "
              "It provides an at-a-glance summary of all your activity.")
    body(doc, "What you see on the dashboard:")
    bullet(doc, "Total Orders — the total count of all orders you have ever submitted.")
    bullet(doc, "Active Orders — orders currently in progress (not yet completed or rejected).")
    bullet(doc, "Waiting for Documents — orders that need you to upload additional files.")
    bullet(doc, "Completed Orders — orders that have been fully processed.")
    bullet(doc, "Recent Orders table — your five most recent orders with quick links.")
    bullet(doc, "Recent Notifications panel — your latest four in-app messages.")
    note(doc, "If the orange 'Waiting for Documents' banner appears, you have orders that "
              "are blocked until you upload the missing files. Act on these first.", "WARNING")

    # 4.2 Browsing
    h2(doc, "4.2  Browsing the Service Catalog")
    body(doc, "The public Services page (/services) lists every active service grouped by category.")
    step(doc, 1, "Click Services in the top navigation bar (visible without login).")
    step(doc, 2, "The page loads a grid of service cards grouped by category.")
    step(doc, 3, "Click on a category tab to filter the view to that category only.")
    step(doc, 4, "Scroll through the service cards. Each card shows the service name, "
                 "category, estimated duration, and base price.")
    step(doc, 5, "Click a service card to view its full details.")

    # 4.3 Searching
    h2(doc, "4.3  Searching & Filtering Services")
    body(doc, "The Services page has a live search bar that filters results as you type.")
    step(doc, 1, "Click inside the Search box at the top of the Services page.")
    step(doc, 2, "Type any keyword — for example: 'passport', 'certificate', 'civil'.")
    step(doc, 3, "Results update automatically after a short pause (350 ms debounce). "
                 "No need to press Enter.")
    step(doc, 4, "To clear the search, delete the text in the search box.")
    step(doc, 5, "Combine search with a category filter by clicking a category tab "
                 "while text is in the search box.")

    # 4.4 Service Detail
    h2(doc, "4.4  Viewing a Service Detail Page")
    body(doc, "Each service has a dedicated detail page showing everything you need to know "
              "before applying.")
    step(doc, 1, "Click a service card on the Services page.")
    step(doc, 2, "The Service Details page loads, showing:")
    bullet(doc, "Service name (bilingual) and category.", 1)
    bullet(doc, "Full description of the service.", 1)
    bullet(doc, "Base price, extra fees, and government fees.", 1)
    bullet(doc, "Estimated duration in working days.", 1)
    bullet(doc, "List of required documents and accepted file types.", 1)
    bullet(doc, "Related services (prerequisites, recommended follow-ups).", 1)
    step(doc, 3, "Click Apply / Create Order to begin the application process.")

    # 4.5 Submitting an Order
    h2(doc, "4.5  Submitting a New Service Request (Order)")
    body(doc, "There are two ways to start a new order:")
    bullet(doc, "From the Service Detail page: click Apply / Create Order.")
    bullet(doc, "From your dashboard: click New Order and then select the service.")
    body(doc, "Step-by-step order submission:")
    step(doc, 1, "Navigate to the service you want and click Apply.")
    step(doc, 2, "If you are not logged in, you will be redirected to the Login page first. "
                 "Log in or register, then return to the service page.")
    step(doc, 3, "The order creation form loads. Fill in any dynamic fields required "
                 "by this specific service (e.g. national ID, date of birth, city).")
    step(doc, 4, "Upload all mandatory documents in the Documents section (see Section 4.6).")
    step(doc, 5, "Review the order summary — service name, price, and document list.")
    step(doc, 6, "Click Submit Order / Confirm.")
    step(doc, 7, "A confirmation screen appears with your unique Order Number "
                 "(format: KH2026-XXXXXX). Write this number down — you can use it to track "
                 "your order without logging in.")
    note(doc, "Submitting an order does not immediately charge you. Payment is handled "
              "separately after the order is processed.", "NOTE")

    # 4.6 Uploading Documents
    h2(doc, "4.6  Uploading Documents to an Order")
    body(doc, "Each service defines which documents are required. The upload component "
              "enforces file type and size rules automatically.")
    step(doc, 1, "On the order creation form (or the order details page), locate the "
                 "Documents section.")
    step(doc, 2, "For each required document, click Choose File or drag and drop a file "
                 "onto the upload area.")
    step(doc, 3, "Only files matching the allowed types (PDF, JPG, PNG, DOCX — as configured "
                 "by the admin) will be accepted. Files over the size limit will be rejected.")
    step(doc, 4, "Once a file is chosen, its name appears next to the upload button.")
    step(doc, 5, "Repeat for each required document.")
    step(doc, 6, "Optional documents can also be uploaded — they are shown separately.")
    step(doc, 7, "Submit the order. All uploaded files are saved securely.")
    note(doc, "If a document rule allows client replacement, you can re-upload a new version "
              "of a document even after submission — as long as an employee has not yet "
              "approved it.", "TIP")

    # 4.7 My Orders
    h2(doc, "4.7  Viewing My Orders List")
    body(doc, "My Orders (/customer/orders) shows all orders you have ever submitted.")
    step(doc, 1, "Click My Orders in the left sidebar of your customer dashboard.")
    step(doc, 2, "The page loads with four summary cards at the top: "
                 "Total, Active, Waiting for Documents, and Completed.")
    step(doc, 3, "The orders table shows: Order Number, Service Name, Status badge, "
                 "Creation Date, and Expected Delivery Date.")
    step(doc, 4, "To search, type in the search box (top of table) — filters by order "
                 "number or service name instantly.")
    step(doc, 5, "If there are more than 15 orders, use the pagination controls at the "
                 "bottom to move between pages.")
    step(doc, 6, "Click View Details on any row to open that order's full detail page.")

    # 4.8 Order Details
    h2(doc, "4.8  Viewing an Order's Details & Timeline")
    body(doc, "The Order Details page shows the complete history and current state of a single order.")
    step(doc, 1, "Click View Details on any order in My Orders.")
    step(doc, 2, "At the top you will see: Order Number, Status badge, Service Name, "
                 "Expected Delivery Date, and any assigned employee or provider.")
    step(doc, 3, "The Documents section shows every file you uploaded, its status "
                 "(Pending, Approved, or Rejected), and when it was uploaded.")
    step(doc, 4, "The Timeline section shows every status change the order has gone through, "
                 "with timestamps and the name of the staff member who made each change.")
    step(doc, 5, "The Notes section shows any messages sent to you by the reviewing employee.")
    step(doc, 6, "If the order is in 'Waiting for Documents' status, you will see a "
                 "Missing Documents section — click the link to respond (see Section 4.10).")
    step(doc, 7, "If the order is completed, a Final Documents section shows the "
                 "deliverable files you can download.")

    # 4.9 Public Tracking
    h2(doc, "4.9  Tracking an Order Without Logging In (Public Tracking)")
    body(doc, "You can check an order's status without logging in using the public "
              "Track Order page.")
    step(doc, 1, "On the public home page, click Track Order (or navigate to /track-order).")
    step(doc, 2, "Enter your Order Number (e.g. KH2026-000001).")
    step(doc, 3, "Enter the Phone number used when you registered.")
    step(doc, 4, "Click Track / Search.")
    step(doc, 5, "A summary of the order status, service name, and expected delivery "
                 "date is displayed. No documents or personal details are shown.")

    # 4.10 Missing Documents Response
    h2(doc, "4.10  Responding to a Missing Documents Request")
    body(doc, "When an employee marks a document as missing, your order status changes to "
              "'Waiting for Customer' and you receive a notification.")
    step(doc, 1, "Click the notification in your notification panel, OR open My Orders and "
                 "find the order with status 'Waiting for Customer'.")
    step(doc, 2, "Open the Order Details page.")
    step(doc, 3, "Click Respond to Missing Documents (or the link in the notification).")
    step(doc, 4, "The Missing Documents Response page shows the list of documents requested "
                 "by the employee, along with any note they wrote.")
    step(doc, 5, "Upload the required file(s) using the Choose File buttons.")
    step(doc, 6, "Add an optional note to the employee if needed.")
    step(doc, 7, "Click Submit Response.")
    step(doc, 8, "The order status automatically moves back to 'Under Review' and the "
                 "employee is notified.")

    # 4.11 Download Final Document
    h2(doc, "4.11  Downloading a Final Document")
    body(doc, "When an order is completed, the provider may upload a final result document "
              "(e.g. a certificate or registration deed). You can download it securely.")
    step(doc, 1, "Open the completed order from My Orders.")
    step(doc, 2, "Scroll to the Final Documents section at the bottom of the page.")
    step(doc, 3, "Click the Download button next to the document.")
    step(doc, 4, "A secure time-limited download token is generated automatically. "
                 "The file downloads to your device.")
    note(doc, "Download links are valid for 30 minutes for security. If the link expires, "
              "simply return to the order page and click Download again.", "NOTE")

    # 4.12 Profile
    h2(doc, "4.12  Managing Your Profile")
    body(doc, "Update your personal information from the Profile page (/customer/profile).")
    step(doc, 1, "Click Profile in the left sidebar, or click your name in the top bar.")
    step(doc, 2, "The Profile page shows your current name, email, and phone number.")
    step(doc, 3, "To update your name or phone: edit the fields and click Save Changes.")
    step(doc, 4, "To change your password: scroll to the Password section, enter your "
                 "current password, then enter and confirm the new password, and click "
                 "Update Password.")
    note(doc, "Your email address is used for login. Contact an administrator if you "
              "need to change it.", "NOTE")

    # 4.13 Notifications
    h2(doc, "4.13  Viewing Notifications")
    step(doc, 1, "Click the Bell icon in the top navigation bar. A red badge shows "
                 "the count of unread notifications.")
    step(doc, 2, "The Notification Panel slides open showing your latest messages, "
                 "each with a title, message body, and timestamp.")
    step(doc, 3, "An unread notification shows a blue dot on the right. "
                 "Reading it marks it as read.")
    step(doc, 4, "Click a notification to navigate directly to the related order.")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 5 — EMPLOYEE
# ═══════════════════════════════════════════════════════════════════════════

def ch5(doc):
    h1(doc, "5.  Employee Guide")
    body(doc, "Employees review customer orders, verify documents, assign providers, "
              "and manage the order lifecycle from submission to completion. "
              "Navigate to your portal at /employee after logging in.")

    h2(doc, "5.1  Employee Dashboard")
    body(doc, "The Employee Dashboard (/employee) shows a real-time summary of your workload.")
    bullet(doc, "Pending Review count — orders waiting for your action.")
    bullet(doc, "Recently assigned orders with quick-access links.")
    bullet(doc, "Latest notifications relevant to your work.")
    body(doc, "Use the sidebar links to navigate: Review Queue, Document Verify, "
              "Missing Service Requests, Reports.")

    h2(doc, "5.2  Review Queue — Viewing & Filtering Orders")
    body(doc, "The Review Queue (/employee/orders) is your primary working list. "
              "It shows all orders assigned to you or awaiting assignment.")
    add_image(doc, "Screenshot (156).png", "Employee / Admin View of Current Orders")
    body(doc, "Summary cards at the top of the queue show:")
    bullet(doc, "Total Orders in queue.")
    bullet(doc, "Under Review — orders currently being processed.")
    bullet(doc, "With Missing Documents — orders that have outstanding document requests.")
    bullet(doc, "Urgent — high-priority orders needing immediate attention.")
    h3(doc, "Available Filters")
    table(doc, [
        ("Search",          "Type an order number or customer name to find specific orders."),
        ("Status",          "Filter by: All, New, Under Review, Waiting Customer, Ready for Delivery, Waiting Government."),
        ("Service",         "Filter by a specific service (dropdown of all active services)."),
        ("Priority",        "Filter by: All, Low, Normal, High, Urgent."),
        ("From Date",       "Show only orders created on or after this date."),
        ("To Date",         "Show only orders created on or before this date."),
        ("Execution Status","Filter by provider status: Assigned, In Progress, Waiting Government, Ready for Delivery."),
        ("Order Ownership", "Show: All, Assigned to Me, Unassigned."),
        ("Missing Docs",    "Show: All, Has Missing Documents, No Missing Documents."),
    ], headers=["Filter", "Description"])
    step(doc, 1, "Use any combination of the filters above to narrow the list.")
    step(doc, 2, "Click Reset Filters to clear all filters and show all orders.")
    step(doc, 3, "Click Review Order on any row to open that order.")

    h2(doc, "5.3  Opening an Order for Review")
    step(doc, 1, "From the Review Queue, click Review Order on the order you want to work on.")
    step(doc, 2, "The Order Review page loads with four metric cards at the top: "
                 "Customer name/phone, Email/National ID, Service name/category, "
                 "and Expected delivery date/city.")
    step(doc, 3, "Below the metrics, the left column shows: Service Requirements & Documents, "
                 "Uploaded Documents list, Notes Log, and Timeline.")
    step(doc, 4, "The right column shows action panels appropriate to the current status: "
                 "Start Review, Missing Documents, Assign Provider, Internal Note, "
                 "Manual Notification, and Reject Order.")

    h2(doc, "5.4  Starting a Review")
    body(doc, "When an order arrives with status 'New', you must start the review before "
              "taking other actions.")
    step(doc, 1, "Open the order from the Review Queue.")
    step(doc, 2, "The amber banner at the top of the right column reads: "
                 "'The order is still in New status. You can start the review now.'")
    step(doc, 3, "Click Start Review.")
    step(doc, 4, "The order status changes to 'Under Review' and the timeline is updated.")
    note(doc, "If you take any other valid action (e.g. request missing documents) while "
              "the order is still 'New', the system automatically moves it to 'Under Review' "
              "first before executing your action.", "TIP")

    h2(doc, "5.5  Verifying & Approving Documents")
    body(doc, "Document verification is done on the dedicated Verify Documents screen.")
    step(doc, 1, "On the Order Review page, scroll to the Uploaded Documents section.")
    step(doc, 2, "Click Open Verification Screen (the blue button in that section).")
    step(doc, 3, "The Employee Verify Documents page opens with a preview pane. "
                 "A secure token is generated automatically for the document preview.")
    step(doc, 4, "Click on a document in the list on the left. Its content (image or PDF) "
                 "appears in the preview pane on the right.")
    step(doc, 5, "Examine the document carefully.")
    step(doc, 6, "Click Approve if the document is valid and complete.")
    step(doc, 7, "Repeat for every required document.")
    note(doc, "The Assign Provider button on the order review page remains disabled until "
              "ALL required documents have been approved. This is enforced by the system.", "WARNING")

    h2(doc, "5.6  Rejecting a Document")
    step(doc, 1, "On the Verify Documents screen, select the document to reject.")
    step(doc, 2, "Click Reject.")
    step(doc, 3, "Enter a clear reason for rejection in the note field.")
    step(doc, 4, "Click Confirm Rejection. The document status changes to 'Rejected'.")
    step(doc, 5, "The order status changes to 'Waiting for Customer' and the customer "
                 "is notified automatically.")

    h2(doc, "5.7  Requesting Missing Documents from the Customer")
    body(doc, "Use this when a required document was not uploaded at all, or when an existing "
              "upload is insufficient but you want the customer to provide a new file.")
    step(doc, 1, "On the Order Review page, scroll to the Request Missing Documents panel "
                 "(right column).")
    step(doc, 2, "Write a clear note explaining what is needed and why.")
    step(doc, 3, "In the checklist below the note, tick each document type the customer "
                 "must provide.")
    step(doc, 4, "Click Send Request.")
    step(doc, 5, "The order status changes to 'Waiting for Customer'. "
                 "A notification is sent to the customer automatically.")
    note(doc, "Be specific in your note. The customer will see exactly what you write "
              "here when they open the Missing Documents Response page.", "TIP")

    h2(doc, "5.8  Adding an Internal Note")
    body(doc, "Internal notes are visible only to employees and admins — not to the customer.")
    step(doc, 1, "Scroll to the Internal Note panel in the right column of the Order Review page.")
    step(doc, 2, "Type your note in the text area.")
    step(doc, 3, "Click Save Note.")
    step(doc, 4, "The note appears in the Notes Log section (left column) under 'Internal Notes'.")

    h2(doc, "5.9  Assigning a Provider")
    body(doc, "After all required documents are approved, you can assign the order to a provider.")
    step(doc, 1, "On the Order Review page, scroll to the Assign Provider panel (right column).")
    step(doc, 2, "If documents are not yet all approved, a warning message explains which "
                 "documents still need verification. The Assign Provider button is disabled.")
    step(doc, 3, "Once all documents are approved, the panel shows a green message: "
                 "'Requirements complete. Choose the appropriate provider.'")
    step(doc, 4, "Review the list of eligible providers shown — each card shows the "
                 "provider's name, type, city, availability status, and approval status.")
    step(doc, 5, "From the Provider dropdown, select the provider you want to assign.")
    step(doc, 6, "Optionally, write an assignment note in the text area.")
    step(doc, 7, "Click Assign Provider.")
    step(doc, 8, "The order status changes to 'Assigned' and the provider is notified.")

    h2(doc, "5.10  Sending a Manual Notification")
    body(doc, "You can send a pre-approved notification template to the customer manually.")
    step(doc, 1, "Scroll to the Manual Approved Notification panel (right column).")
    step(doc, 2, "Select a notification template from the dropdown list.")
    step(doc, 3, "Click Send Notification.")
    note(doc, "Only pre-approved templates can be sent from this screen. "
              "You cannot write a free-form message here.", "NOTE")

    h2(doc, "5.11  Returning an Order to the Provider")
    body(doc, "If the provider submits a result that is insufficient, you can send it back.")
    step(doc, 1, "The order must be in 'Ready for Delivery' status (provider submitted result).")
    step(doc, 2, "Scroll to the Review Provider Result section.")
    step(doc, 3, "Click the Return to Provider tab or button.")
    step(doc, 4, "Write a clear reason in the text area explaining what needs to be corrected.")
    step(doc, 5, "Click Return to Provider.")
    step(doc, 6, "The order status changes to 'In Progress' and the provider is notified.")

    h2(doc, "5.12  Returning an Order to Internal Review")
    body(doc, "If a completed result needs internal re-examination before you accept it:")
    step(doc, 1, "The order must be in 'Ready for Delivery' status.")
    step(doc, 2, "Scroll to the Review Provider Result section.")
    step(doc, 3, "Write a reason in the Return to Internal Review text area.")
    step(doc, 4, "Click Return to Internal Review.")
    step(doc, 5, "The order status changes back to 'Under Review'.")

    h2(doc, "5.13  Completing (Approving) an Order")
    body(doc, "When the provider's final result is satisfactory, you complete the order.")
    step(doc, 1, "The order must be in 'Ready for Delivery' status and a final document "
                 "must have been uploaded by the provider.")
    step(doc, 2, "Scroll to the Review Provider Result section. The final document is "
                 "displayed with a preview.")
    step(doc, 3, "Verify the final document is correct.")
    step(doc, 4, "Click Approve Result and Complete Order.")
    step(doc, 5, "The order status changes to 'Completed'. The customer is notified "
                 "and can download the final document.")

    h2(doc, "5.14  Rejecting an Order")
    body(doc, "Reject an order when it cannot be fulfilled (e.g. fraudulent documents, "
              "ineligible applicant).")
    step(doc, 1, "Scroll to the bottom of the right column on the Order Review page.")
    step(doc, 2, "Click the red Reject Order button.")
    step(doc, 3, "A confirmation modal appears: 'Confirm order rejection — this action "
                 "cannot be undone.'")
    step(doc, 4, "Click Yes, Reject Order to confirm.")
    step(doc, 5, "The order status changes to 'Rejected', the customer is notified, "
                 "and the rejection reason is logged in the timeline.")
    note(doc, "Rejection is permanent and cannot be undone. Use this only when necessary.", "WARNING")

    h2(doc, "5.15  Employee Reports")
    body(doc, "The Reports page (/employee/reports) provides operational analytics.")
    step(doc, 1, "Click Reports in the left sidebar.")
    step(doc, 2, "Set the desired date range using the From and To date pickers.")
    step(doc, 3, "The report shows: orders processed, average processing time per service, "
                 "document rejection rates, and volume by service.")
    step(doc, 4, "Use the export function to download the report as a file if available.")

    h2(doc, "5.16  Missing Service Requests")
    body(doc, "Customers can submit requests for services that do not exist yet. "
              "Employees can view these to pass them on to admins.")
    step(doc, 1, "Click Missing Service Requests in the left sidebar.")
    step(doc, 2, "The table shows all requests: customer name, requested service description, "
                 "date submitted, and status.")
    step(doc, 3, "Review the requests and communicate significant ones to the admin team.")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 6 — PROVIDER
# ═══════════════════════════════════════════════════════════════════════════

def ch6(doc):
    h1(doc, "6.  Provider Guide")
    body(doc, "Providers execute the actual service after employee review is complete. "
              "They receive assigned orders, process them, and upload the final result.")

    h2(doc, "6.1  Provider Dashboard")
    body(doc, "The Provider Dashboard (/provider) shows:")
    bullet(doc, "Count of newly assigned orders.")
    bullet(doc, "Orders currently in progress.")
    bullet(doc, "Recent completions.")
    bullet(doc, "Latest notifications.")

    h2(doc, "6.2  Viewing Assigned Orders")
    step(doc, 1, "Click Assigned Orders in the left sidebar (/provider/orders).")
    step(doc, 2, "The table shows all orders assigned to you: Order Number, Customer, "
                 "Service, Status, City, and Date.")
    step(doc, 3, "Use the available filters to narrow by status or date range.")
    step(doc, 4, "Click View Details on any order to open it.")

    h2(doc, "6.3  Opening an Assigned Order")
    step(doc, 1, "Click View Details on an order row.")
    step(doc, 2, "The Provider Order Details page shows:")
    bullet(doc, "Order header: number, status, service name.", 1)
    bullet(doc, "Customer information: name and phone.", 1)
    bullet(doc, "All verified documents the employee approved.", 1)
    bullet(doc, "Order timeline.", 1)
    bullet(doc, "Status update panel and file upload area.", 1)

    h2(doc, "6.4  Updating Order Status")
    step(doc, 1, "On the Provider Order Details page, locate the Status Update panel.")
    step(doc, 2, "Select the new status from the dropdown:")
    bullet(doc, "In Progress — you have started working on the order.", 1)
    bullet(doc, "Waiting for External Party — the order requires a government entity's involvement.", 1)
    bullet(doc, "Ready for Delivery — the service is complete and result is ready.", 1)
    step(doc, 3, "Add a note explaining the update (optional but recommended).")
    step(doc, 4, "Click Update Status.")
    step(doc, 5, "The employee is notified of the status change.")

    h2(doc, "6.5  Uploading the Final Deliverable Document")
    body(doc, "When the service is complete, upload the result document so the employee "
              "can review it and the customer can download it.")
    step(doc, 1, "On the Provider Order Details page, locate the Upload Result section.")
    step(doc, 2, "Click Choose File and select the final document (PDF, JPG, or other "
                 "supported format).")
    step(doc, 3, "Mark the upload as 'Final Document' if the toggle is available.")
    step(doc, 4, "Add any relevant notes.")
    step(doc, 5, "Click Upload.")
    step(doc, 6, "Set the order status to 'Ready for Delivery'.")
    step(doc, 7, "The reviewing employee is notified and will either approve or return the result.")
    note(doc, "You can only see documents that the admin has configured as 'Provider Can View'. "
              "Customer personal data not related to the service is not shown.", "NOTE")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 7 — ADMIN
# ═══════════════════════════════════════════════════════════════════════════

def ch7(doc):
    h1(doc, "7.  Administrator Guide")
    body(doc, "Administrators have full access to configure every aspect of Khalisni. "
              "The admin portal is accessible at /admin after logging in with an admin account. "
              "The left sidebar provides links to every management section.")

    # 7.1 Overview
    h2(doc, "7.1  Admin Dashboard Overview")
    body(doc, "URL: /admin — The dashboard gives a high-level operational snapshot.")
    add_image(doc, "Screenshot (156).png", "Admin Dashboard — Operational Overview")
    body(doc, "Six metric cards at the top show:")
    bullet(doc, "New Orders Today — orders created since midnight.")
    bullet(doc, "In Progress — orders actively being worked on.")
    bullet(doc, "Waiting for Customer — orders blocked on customer action.")
    bullet(doc, "Completed This Week — successfully closed orders this week.")
    bullet(doc, "Delayed Orders — orders past their expected delivery date.")
    bullet(doc, "Revenue Estimate — sum of base prices of all completed orders.")
    body(doc, "Two charts below the metrics show:")
    bullet(doc, "Orders by Status — bar chart of order counts per status.")
    bullet(doc, "Most Requested Services — pie chart of order volume per service.")

    # 7.2 Orders Management
    h2(doc, "7.2  Orders Management")
    body(doc, "URL: /admin/orders — Admins can view all orders system-wide with full filters.")
    step(doc, 1, "Click Orders Management (إدارة الطلبات) in the left sidebar.")
    step(doc, 2, "The table shows all orders across all customers, employees, and providers.")
    step(doc, 3, "Use the filter panel to narrow by status, service, employee, provider, "
                 "date range, priority, or city.")
    step(doc, 4, "Click View Details on any order to open the Admin Order Details page.")
    body(doc, "The Admin Order Details page provides the same information as the Employee "
              "review page, plus admin-only actions such as forcibly updating status or "
              "overriding assignments.")

    # 7.3 Rule Management
    h2(doc, "7.3  Rule Management Hub")
    body(doc, "URL: /admin/rules — This is the central configuration hub for all business rules. "
              "Access it via Rules (قواعد التشغيل) in the sidebar.")
    add_image(doc, "Screenshot (157).png", "Rule Management — Main Tab Navigation")
    body(doc, "The page has tabs across the top:")
    bullet(doc, "Notification Templates   |   Workflow Rules   |   Provider Rules")
    bullet(doc, "Required Documents   |   Services and Pricing   |   System Settings")
    bullet(doc, "Audit Logs   |   Users and Roles   |   Payments")
    note(doc, "All changes made in Rule Management are recorded in the Audit Log automatically.", "NOTE")

    # 7.3.1
    h3(doc, "7.3.1  Services & Pricing — Viewing the List")
    body(doc, "Click the Services and Pricing tab.")
    add_image(doc, "Screenshot (158).png", "Services & Pricing — Service List Table")
    body(doc, "The table shows all services. Columns:")
    table(doc, [
        ("Service",            "Bilingual service name."),
        ("Category",           "Parent category for grouping in the catalog."),
        ("Base Price",         "Main fee in JOD shown to the customer."),
        ("Estimated Duration", "Expected working days for completion."),
        ("Employee Review",    "Yes = employee review required before provider assignment."),
        ("Provider Required",  "Yes = order must be assigned to a provider."),
        ("Status",             "Active (shown to customers) or Inactive (hidden)."),
        ("Actions",            "Edit and Disable buttons."),
    ], headers=["Column", "Description"])

    h3(doc, "7.3.2  Services & Pricing — Creating a New Service")
    step(doc, 1, "Click the New Service button (top right of the table).")
    add_image(doc, "Screenshot (159).png", "New Service Form — Basic Details (Part 1)")
    add_image(doc, "Screenshot (160).png", "New Service Form — Fees & Duration (Part 2)")
    add_image(doc, "Screenshot (161).png", "New Service Form — Flags & Activation (Part 3)")
    step(doc, 2, "Fill in the form fields:")
    table(doc, [
        ("Arabic name",            "Service name in Arabic — shown to Arabic-language users."),
        ("English name",           "Service name in English — shown in English mode and reports."),
        ("Category",               "Select the parent category from the dropdown."),
        ("Base price (JOD)",       "The main service fee charged to the customer. Use 0.00 for free services."),
        ("Extra service fee",      "Optional internal platform surcharge on top of the base price."),
        ("Government fee",         "Optional government-imposed charge to display separately."),
        ("Duration unit",          "Days (default). Used to calculate expected delivery date."),
        ("Estimated duration",     "Integer number of duration units. Shown to customer on the service page."),
        ("Arabic description",     "Full bilingual service description displayed on the service detail page."),
        ("English description",    "English version of the description."),
        ("Provider is required",   "Toggle ON if a provider must execute this service."),
        ("Employee review is req.", "Toggle ON if an employee must review and verify documents."),
        ("Service is active",      "Toggle OFF to hide the service from the public catalog without deleting it."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Save Service.")
    step(doc, 4, "The new service appears in the list and is immediately visible in the "
                 "public catalog (if active).")
    note(doc, "Price changes take effect immediately for all new orders. Existing orders "
              "keep the price that was set when they were created.", "WARNING")

    h3(doc, "7.3.3  Services & Pricing — Editing a Service")
    step(doc, 1, "Find the service in the list and click Edit.")
    step(doc, 2, "The same form as New Service appears with existing values pre-filled.")
    step(doc, 3, "Modify any field and click Save Service.")
    step(doc, 4, "All changes are logged in the Audit Log.")

    h3(doc, "7.3.4  Services & Pricing — Disabling a Service")
    step(doc, 1, "Find the service in the list and click Disable.")
    step(doc, 2, "A confirmation prompt appears. Click Confirm.")
    step(doc, 3, "The service status changes to Inactive. It disappears from the public "
                 "catalog but all existing orders for it are unaffected.")
    note(doc, "Disabling is reversible. To re-enable, click Edit on the service and "
              "toggle 'Service is active' back ON.", "TIP")

    h3(doc, "7.3.5  Required Documents — Creating a Document Rule")
    body(doc, "Click the Required Documents tab.")
    step(doc, 1, "Click New Document Rule.")
    add_image(doc, "Screenshot (162).png", "New Document Rule Form — Document Identity")
    add_image(doc, "Screenshot (163).png", "New Document Rule — File Types & Max Size")
    add_image(doc, "Screenshot (164).png", "New Document Rule — Behaviour Flags")
    step(doc, 2, "Fill in the form:")
    table(doc, [
        ("Service",                "The service this document rule applies to. Select from dropdown."),
        ("Arabic document name",   "Document label shown to Arabic-speaking users."),
        ("English document name",  "Document label shown to English-speaking users."),
        ("Accepted file types",    "Tick all allowed formats: JPG, PNG, PDF, DOCX, DOC, etc."),
        ("Maximum size (bytes)",   "File size limit. Default: 10,485,760 bytes (10 MB)."),
        ("Required document",      "ON = order cannot progress to review without this file being uploaded and approved."),
        ("Requires verification",  "ON = an employee must manually approve this document."),
        ("Client can replace",     "ON = the customer can re-upload a new version even after submission."),
        ("Provider can view",      "ON = the assigned provider can see this document in their portal."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Save Document Rule.")

    h3(doc, "7.3.6  Provider Assignment Rules")
    body(doc, "Click the Provider Rules tab. Assignment rules automatically route approved "
              "orders to the correct provider.")
    step(doc, 1, "Click New Assignment Rule.")
    add_image(doc, "Screenshot (165).png", "New Provider Assignment Rule Form")
    step(doc, 2, "Select the Service this rule applies to.")
    step(doc, 3, "Select the Provider to assign.")
    step(doc, 4, "Ensure Assignment rule is active is checked.")
    step(doc, 5, "Click Save Assignment Rule.")
    note(doc, "Only approved, active providers appear in the Provider dropdown.", "NOTE")

    h3(doc, "7.3.7  Workflow Rules")
    body(doc, "Click the Workflow Rules tab. Workflow rules define automated actions that "
              "trigger at specific points in the order lifecycle — for example, "
              "auto-approving a document type or auto-routing to a provider when conditions are met.")
    body(doc, "To create a workflow rule:")
    step(doc, 1, "Click New Workflow Rule.")
    step(doc, 2, "Define the trigger condition (e.g. 'When order status becomes UNDER_REVIEW').")
    step(doc, 3, "Define the action (e.g. 'Auto-approve document of type X').")
    step(doc, 4, "Set any additional conditions (service, priority, etc.).")
    step(doc, 5, "Toggle the rule Active and click Save.")

    h3(doc, "7.3.8  Notification Templates — Creating a Template")
    body(doc, "Click the Notification Templates tab.")
    add_image(doc, "Screenshot (166).png", "Notification Templates — Template List")
    step(doc, 1, "Click New Template.")
    add_image(doc, "Screenshot (167).png", "New Notification Template — Part 1")
    add_image(doc, "Screenshot (168).png", "New Notification Template — Part 2")
    add_image(doc, "Screenshot (169).png", "New Notification Template — Activation")
    step(doc, 2, "Fill in the form:")
    table(doc, [
        ("Template key",    "Unique internal identifier (short, lowercase, use underscores). "
                            "Example: order_approved_customer"),
        ("Channel",         "System (in-app) is the currently supported channel."),
        ("Arabic title",    "Short notification heading shown to Arabic users."),
        ("Arabic message",  "Notification body in Arabic. Use placeholders: "
                            "{order_number}, {service_name}, {customer_name}, {status_label}."),
        ("English title",   "Short notification heading in English."),
        ("English message", "Notification body in English with the same placeholder support."),
        ("Template is active", "Toggle ON to activate. Inactive templates are skipped."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Preview to see how the message will look before saving.")
    step(doc, 4, "Click Save Template.")
    note(doc, "Always preview a template before activating it to verify that placeholders "
              "are correctly named.", "WARNING")

    h3(doc, "7.3.9  Payments Management")
    body(doc, "Click the Payments tab.")
    add_image(doc, "Screenshot (170).png", "Payments — Payment Status List")
    body(doc, "This tab shows all payment records linked to orders.")
    bullet(doc, "Payment status can only be updated through the safe status-action workflow.")
    bullet(doc, "Raw transaction records cannot be edited directly — this is by design "
               "to maintain a complete audit trail.")
    body(doc, "To update a payment status:")
    step(doc, 1, "Find the payment record in the list using the search or date filters.")
    step(doc, 2, "Click Update Status next to the record.")
    step(doc, 3, "Select the new status (e.g. Paid, Pending, Refunded).")
    step(doc, 4, "Add a note explaining the change.")
    step(doc, 5, "Click Save.")

    h3(doc, "7.3.10  System Settings")
    body(doc, "Click the System Settings tab. Only whitelisted, safe settings can be edited here.")
    add_image(doc, "Screenshot (176).png", "New System Setting Form (English)")
    add_image(doc, "Screenshot (178).png", "New System Setting Form (Arabic)")
    step(doc, 1, "Click New Setting.")
    step(doc, 2, "Select the setting key from the dropdown (e.g. site_homepage, main_title).")
    step(doc, 3, "Fill in the value fields:")
    table(doc, [
        ("Setting",        "Pre-approved setting key chosen from the allowed list."),
        ("Homepage content","The public-facing text or value for the home page."),
        ("Main title",     "Short headline displayed at the top of the home page."),
        ("Supporting text","Subtitle or supporting description below the main title."),
        ("Admin note",     "Internal note explaining why this setting is being changed (not public)."),
    ], headers=["Field", "Description"])
    step(doc, 4, "Click Save Setting.")

    h3(doc, "7.3.11  Audit Logs")
    body(doc, "Click the Audit Logs tab.")
    add_image(doc, "Screenshot (174).png", "Audit Log — Filter Panel")
    add_image(doc, "Screenshot (175).png", "Audit Log — Results Table")
    body(doc, "The Audit Log is a read-only record of every significant action in the system.")
    body(doc, "How to search the audit log:")
    step(doc, 1, "Set the From Date and To Date to define the search window.")
    step(doc, 2, "Optionally filter by Action (e.g. create_service_category, logout, register).")
    step(doc, 3, "Optionally filter by Module (e.g. Service, Payment, Notification, Order).")
    step(doc, 4, "Optionally filter by User (type a name or ID).")
    step(doc, 5, "Select Result: All Results / Success / Failed.")
    step(doc, 6, "Click Refresh Logs.")
    step(doc, 7, "The results table shows: Date/Time, Message, Result, Action, Module, User.")
    note(doc, "Audit logs cannot be edited or deleted. They are the authoritative record "
              "for compliance reviews.", "NOTE")

    h3(doc, "7.3.12  Users & Roles — Creating a User")
    body(doc, "Click the Users and Roles tab.")
    add_image(doc, "Screenshot (171).png", "Users & Roles — User List")
    step(doc, 1, "Click New User.")
    add_image(doc, "Screenshot (172).png", "New User Form — Identity & Credentials")
    add_image(doc, "Screenshot (173).png", "New User Form — Role & Status")
    step(doc, 2, "Fill in all required fields:")
    table(doc, [
        ("Full name",        "User's display name (business-facing)."),
        ("Email",            "Login email. Must be unique in the system."),
        ("Phone",            "Primary contact number for notifications."),
        ("Password",         "Temporary initial password. The user should change this on first login."),
        ("Role",             "Customer / Employee / Provider / Support. Leave blank if not applicable."),
        ("National ID",      "Optional identity reference number."),
        ("User is active",   "Toggle ON to allow login. Toggle OFF to suspend access without deleting."),
        ("User is verified", "Toggle ON once you have completed business verification for this user."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Save User.")
    step(doc, 4, "The user can now log in with their email and the initial password.")
    note(doc, "Remind new users to change their password immediately after first login.", "TIP")

    h3(doc, "7.3.13  Users & Roles — Editing a User's Permissions")
    step(doc, 1, "Find the user in the Users & Roles list and click Edit.")
    step(doc, 2, "The user detail panel opens with two tabs: Data and Permissions.")
    step(doc, 3, "Click the Permissions tab.")
    step(doc, 4, "All available Django permissions are shown grouped by app/module.")
    step(doc, 5, "Check or uncheck individual permissions as needed.")
    step(doc, 6, "Click Save Permissions.")
    note(doc, "Normal admin accounts cannot grant super-admin level permissions. "
              "Super-admin accounts are protected and cannot be edited from this screen.", "WARNING")

    # 7.4 Service Categories
    h2(doc, "7.4  Service Categories")
    body(doc, "URL: /admin/service-categories — Organise the public catalog into a browsable hierarchy.")
    add_image(doc, "Screenshot (180).png", "Service Category Management — List View")
    step(doc, 1, "Click New Category.")
    add_image(doc, "Screenshot (181).png", "New Service Category — Part 1")
    add_image(doc, "Screenshot (182).png", "New Service Category — Part 2")
    step(doc, 2, "Fill in the form:")
    table(doc, [
        ("English name",       "Category label for English users."),
        ("Arabic name",        "Category label for Arabic users."),
        ("Slug",               "URL-safe identifier (auto-generated — leave blank to derive from name)."),
        ("Parent category",    "Assign a parent to make this a sub-category. Leave blank for root."),
        ("Root category",      "Check if this is a top-level category."),
        ("Color",              "Hex color code for the category tile (e.g. #4F9FDE)."),
        ("Icon",               "FontAwesome icon class (e.g. fa-passport). Used on the catalog tile."),
        ("Description",        "Short optional description shown on the catalog page."),
        ("Sort order",         "Integer controlling display order (lower number = displayed first)."),
        ("Show on public site","Toggle ON to display in the public catalog. OFF = internal-only."),
        ("Category is active", "Toggle OFF to hide the category entirely without deleting it."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Create Category.")
    body(doc, "To edit an existing category:")
    step(doc, 1, "Find it in the list and click Edit.")
    step(doc, 2, "Modify the fields and click Save.")
    body(doc, "To delete a category:")
    step(doc, 1, "Click Delete on a category that has no associated services.")
    note(doc, "A category with active services cannot be deleted. Disable the services first.", "WARNING")

    # 7.5 Service Relations
    h2(doc, "7.5  Service Relations")
    body(doc, "URL: /admin/service-relations — Define business dependencies and "
              "recommendations between services.")
    add_image(doc, "Screenshot (186).png", "Service Relations — List View")
    step(doc, 1, "Click New Relation.")
    add_image(doc, "Screenshot (187).png", "New Service Relation — Part 1")
    add_image(doc, "Screenshot (188).png", "New Service Relation — Part 2")
    step(doc, 2, "Fill in the form:")
    table(doc, [
        ("Source service",   "The service that triggers the relation (the 'starting' service)."),
        ("Target service",   "The service affected by or related to the source."),
        ("Source category",  "Optional filter to narrow the source service dropdown."),
        ("Target category",  "Optional filter to narrow the target service dropdown."),
        ("Rule type",        "Choose the relationship meaning (see table below)."),
        ("Customer message", "A plain-language note shown to customers on the service detail page."),
        ("This prerequisite is mandatory", "ON = customer is blocked from ordering target until source is completed."),
        ("Relation is active", "Toggle OFF to suspend the relation without deleting it."),
    ], headers=["Field", "Description"])
    table(doc, [
        ("Required Before",    "The source must be completed before the customer can apply for the target."),
        ("Recommended After",  "The system recommends the target after the source is completed."),
        ("Optional Bundle",    "The target is offered as an optional add-on when applying for the source."),
        ("Alternative Option", "The target is presented as an alternative path to the source."),
    ], headers=["Rule Type", "Behaviour"])
    step(doc, 3, "Click Create Relation.")

    # 7.6 Providers
    h2(doc, "7.6  Providers Management — Creating a Provider")
    body(doc, "URL: /admin/providers — Create and manage provider accounts.")
    add_image(doc, "Screenshot (206).png", "Providers Management — Provider List")
    step(doc, 1, "Click New Provider.")
    add_image(doc, "Screenshot (207).png", "New Provider Form — Part 1")
    add_image(doc, "Screenshot (208).png", "New Provider Form — Part 2 (Services & Flags)")
    step(doc, 2, "Fill in the form:")
    table(doc, [
        ("Full name",         "Provider's display name (individual or organisation)."),
        ("Email",             "Login email for the provider account."),
        ("Password",          "Temporary initial password."),
        ("Provider type",     "Individual or Organisation."),
        ("Commercial name",   "Official business / trading name."),
        ("Internal code",     "Optional internal reference number."),
        ("Address",           "Physical address of the provider."),
        ("Tax number",        "VAT / tax registration number."),
        ("Assigned services", "Multi-select the services this provider is qualified to execute."),
        ("Account is active", "Toggle OFF to suspend the provider without deleting."),
        ("Verified",          "Toggle ON once the provider has been onboarded and verified."),
        ("Accepts payments",  "Toggle ON if this provider handles payment collection for orders."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Add Provider.")
    note(doc, "The provider can now log in at /provider with their email and initial password.", "TIP")

    # 7.7 Provider Service Assignments
    h2(doc, "7.7  Provider Service Assignments")
    body(doc, "URL: /admin/provider-services — Link providers to specific services and cities.")
    add_image(doc, "Screenshot (209).png", "Assign Service to Provider Form")
    step(doc, 1, "Click New Relation (ربط جديد).")
    step(doc, 2, "Select the Provider from the dropdown.")
    step(doc, 3, "Select the City.")
    step(doc, 4, "Select the Service to assign.")
    step(doc, 5, "Ensure Link is active is checked.")
    step(doc, 6, "Click Add Link (ربط الخدمة).")

    # 7.8 Missing Service Requests
    h2(doc, "7.8  Missing Service Requests (Admin)")
    body(doc, "URL: /admin/missing-service-requests — View requests from customers for services "
              "that do not yet exist in the catalog.")
    add_image(doc, "Screenshot (202).png", "Missing Service Requests — Admin View")
    step(doc, 1, "Click Missing Service Requests in the left sidebar.")
    step(doc, 2, "The page shows all submitted requests with filters for status and date.")
    step(doc, 3, "Review the requests to identify popular demand for new services.")
    step(doc, 4, "Create the new service in Services & Pricing when ready.")

    # 7.9 Public Site Management
    h2(doc, "7.9  Public Site Management")
    body(doc, "URL: /admin/public-site — Control the entire customer-facing website "
              "without any development work.")
    add_image(doc, "Screenshot (189).png", "Public Site Management — Module Overview")
    body(doc, "Four sub-modules are available:")
    bullet(doc, "Homepage Content Editor — edit all text and buttons on the home page.")
    bullet(doc, "Advertisement Manager — create and schedule promotional banners.")
    bullet(doc, "Theme Settings — customise colours, logo, and fonts.")
    bullet(doc, "Preview Public Page — see exactly how the site looks to visitors.")

    h3(doc, "7.9.1  Homepage Content Editor")
    body(doc, "URL: /admin/public-site/content")
    add_image(doc, "Screenshot (190).png", "Homepage Content Editor — Hero & How It Works")
    add_image(doc, "Screenshot (191).png", "Homepage Content Editor — Hero Buttons")
    add_image(doc, "Screenshot (192).png", "Homepage Content Editor — Contact & Footer")
    add_image(doc, "Screenshot (193).png", "Homepage Content Editor — Active Toggle")
    body(doc, "To create or edit homepage content:")
    step(doc, 1, "Click New Version (or click Edit on an existing version).")
    step(doc, 2, "Fill in all bilingual fields:")
    table(doc, [
        ("Version name",          "Internal label for this content set (e.g. 'June 2026 Launch')."),
        ("Hero title AR / EN",    "Large heading text on the hero banner, in Arabic and English."),
        ("Hero subtitle AR / EN", "Smaller supporting text under the hero title."),
        ("Primary button text",   "Label on the main CTA button (e.g. 'Get Started' / 'ابدأ الآن')."),
        ("Primary button URL",    "Destination URL (e.g. /create-order/)."),
        ("Secondary button text", "Label on the secondary button (e.g. 'Track your order')."),
        ("Secondary button URL",  "Destination URL (e.g. /track-order/)."),
        ("Phone / WhatsApp",      "Contact phone numbers shown in the contact section and footer."),
        ("Email address",         "Contact email shown in the footer."),
        ("Office address AR/EN",  "Physical address in both languages."),
        ("Footer text AR/EN",     "Copyright and tagline text at the bottom of every page."),
        ("Hero image",            "Upload a banner background image (JPG or PNG, max 10 MB)."),
        ("Active content version","Toggle ON to publish this version. Only one version is active at a time."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Save Content.")
    step(doc, 4, "To make this version live, toggle Active Content Version ON and click Save again.")
    note(doc, "Only one content version can be active at a time. Activating a new version "
              "automatically deactivates the previous one.", "WARNING")

    h3(doc, "7.9.2  Advertisement Manager")
    body(doc, "URL: /admin/public-site/advertisements — Create and schedule banners or "
              "announcements shown on the public site.")
    add_image(doc, "Screenshot (194).png", "Advertisement Manager — List View")
    step(doc, 1, "Click New Advertisement (إعلان جديد).")
    add_image(doc, "Screenshot (195).png", "New Advertisement — Title & Description")
    add_image(doc, "Screenshot (196).png", "New Advertisement — Type & Image")
    add_image(doc, "Screenshot (197).png", "New Advertisement — Button & Colors")
    add_image(doc, "Screenshot (198).png", "New Advertisement — Dates")
    add_image(doc, "Screenshot (199).png", "New Advertisement — Final Review & Active Toggle")
    step(doc, 2, "Fill in all fields:")
    table(doc, [
        ("Title AR",          "Advertisement heading in Arabic."),
        ("Title EN",          "Advertisement heading in English."),
        ("Description AR",    "Body text in Arabic."),
        ("Description EN",    "Body text in English."),
        ("Type",              "General, Banner, or Announcement. Controls display style."),
        ("Image",             "Optional banner background or illustration image."),
        ("Button AR / EN",    "Call-to-action button label in each language."),
        ("Button URL",        "URL the button links to."),
        ("Text color",        "Hex color for all overlay text."),
        ("Background color",  "Hex color for the banner background."),
        ("Display order",     "Integer — lower numbers appear first."),
        ("Start date",        "Date the ad becomes visible."),
        ("End date",          "Date the ad is automatically hidden."),
        ("Active",            "Manual on/off override — takes precedence over date range."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Add Advertisement (إضافة الإعلان).")
    body(doc, "To edit or deactivate an advertisement:")
    step(doc, 1, "Find it in the list and click Edit (تعديل).")
    step(doc, 2, "Change any field or toggle Active to OFF.")
    step(doc, 3, "Click Save.")

    h3(doc, "7.9.3  Theme Settings")
    body(doc, "URL: /admin/public-site/theme — Customise the visual identity of the site.")
    add_image(doc, "Screenshot (200).png", "Theme Settings — Configuration")
    add_image(doc, "Screenshot (201).png", "Theme Settings — Live Preview")
    step(doc, 1, "Click Theme Settings in the Public Site Management menu.")
    step(doc, 2, "Configure the fields:")
    table(doc, [
        ("Theme name",        "Internal label (e.g. 'Default Khalisni Theme')."),
        ("Logo",              "Upload site logo (PNG with transparent background recommended)."),
        ("Browser tab icon",  "Favicon file (16x16 or 32x32 ICO or PNG)."),
        ("Primary color",     "Main brand color — used for buttons, links, and highlights."),
        ("Secondary color",   "Accent color for secondary buttons and badges."),
        ("Text color",        "Default body text color."),
        ("Page background",   "Main page background color."),
        ("Footer background", "Footer section background color."),
        ("Header background", "Navigation bar background color."),
        ("Active theme",      "Toggle ON to apply this theme to the live site immediately."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Watch the live Header Preview and Footer Preview panels on the right "
                 "to see the effect of your color choices in real time.")
    step(doc, 4, "Click Save Theme.")
    note(doc, "Only one theme can be active at a time.", "NOTE")

    h3(doc, "7.9.4  Preview Public Page")
    body(doc, "URL: /admin/public-site/preview")
    step(doc, 1, "Click Preview Public Page from the Public Site Management hub.")
    step(doc, 2, "A full read-only preview of the live public home page is displayed "
                 "within the admin panel.")
    step(doc, 3, "Use this to verify how content and theme changes look before announcing them.")

    # 7.10 Help Guide
    h2(doc, "7.10  Help Guide Management")
    body(doc, "URL: /admin/help-guides — Create contextual in-app guide entries that appear "
              "as help panels next to specific screens or actions.")
    add_image(doc, "Screenshot (210).png", "Help Guide — New Screenshot Entry Form")
    step(doc, 1, "Click New Entry or New Screenshot.")
    step(doc, 2, "Fill in the form:")
    table(doc, [
        ("Guide page",       "Select which screen or workflow step this guide entry belongs to."),
        ("Caption",          "Short label for the screenshot thumbnail."),
        ("Preview summary",  "Brief text shown when the help panel is collapsed."),
        ("Upload image",     "The actual screenshot to embed in the guide panel."),
        ("Static path",      "Alternative: use a static file path instead of an uploaded image."),
        ("Placeholder label","Text displayed if the image cannot be loaded."),
    ], headers=["Field", "Description"])
    step(doc, 3, "Click Create Entry.")
    note(doc, "Help guide entries appear as blue help icons (?) next to the relevant "
              "actions on the employee review page. Employees can click them to see "
              "guidance without leaving their current screen.", "TIP")

    # 7.11 Reports
    h2(doc, "7.11  Admin Reports")
    body(doc, "URL: /admin/reports — Full platform analytics available only to admins.")
    step(doc, 1, "Click Reports (التقارير) in the left sidebar.")
    step(doc, 2, "Set the date range filters.")
    step(doc, 3, "The report displays: total orders by period, revenue by service, "
                 "employee performance, provider turnaround times, and rejection rates.")
    step(doc, 4, "Export the report using the download button.")

    # 7.12 Notifications
    h2(doc, "7.12  Notifications Management")
    body(doc, "URL: /admin/notifications — View all system-wide notifications dispatched.")
    step(doc, 1, "Click Notifications (الإشعارات) in the left sidebar.")
    step(doc, 2, "Filter by user, date range, or notification type.")
    step(doc, 3, "The table shows: Notification title, recipient, channel, sent date, "
                 "and read status.")
    note(doc, "This page is for monitoring only. To create new notification templates, "
              "use the Notification Templates tab in Rule Management.", "NOTE")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 8 — ORDER STATUS REFERENCE
# ═══════════════════════════════════════════════════════════════════════════

def ch8(doc):
    h1(doc, "8.  Order Status Reference")
    body(doc, "Every order moves through a defined set of statuses. Understanding these "
              "statuses helps all users know what is expected at each stage.")
    status_table(doc, [
        ("NEW",                 "New",                    "Order submitted by customer. Awaiting employee review."),
        ("UNDER_REVIEW",        "Under Review",           "Employee has opened the order and is reviewing documents."),
        ("WAITING_CUSTOMER",    "Waiting for Customer",   "Employee has requested additional documents. Customer must respond."),
        ("ASSIGNED",            "Assigned to Provider",   "All documents approved. Order sent to provider for execution."),
        ("IN_PROGRESS",         "In Progress",            "Provider is actively working on the order."),
        ("WAITING_GOVERNMENT",  "Waiting — External Party","Order requires action from a third-party government entity."),
        ("READY_FOR_DELIVERY",  "Ready for Delivery",     "Provider has uploaded the result. Employee must review and approve."),
        ("COMPLETED",           "Completed",              "Order fully processed. Final document available for download."),
        ("CANCELLED",           "Cancelled",              "Order was cancelled before completion."),
        ("REJECTED",            "Rejected",               "Order rejected by employee. Cannot be resubmitted."),
    ])
    body(doc, "Priority levels that can be assigned to orders:")
    table(doc, [
        ("LOW",    "Below normal urgency. Processed in sequence."),
        ("NORMAL", "Standard priority. Default for all new orders."),
        ("HIGH",   "Elevated urgency. Should be reviewed before NORMAL orders."),
        ("URGENT", "Highest priority. Requires immediate action."),
    ], headers=["Priority", "Meaning"])
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 9 — FAQ
# ═══════════════════════════════════════════════════════════════════════════

def ch9(doc):
    h1(doc, "9.  Frequently Asked Questions")
    faqs = [
        ("Can I submit an order without creating an account?",
         "Order submission requires a registered account. However, you can track any existing order "
         "publicly using its Order Number and the phone number registered on the account — no login needed."),
        ("How long will my order take?",
         "Each service page shows an Estimated Duration in working days. This is a target, not a guarantee. "
         "Complex orders or missing documents can cause delays. You can check the real-time status at any time."),
        ("What file formats are accepted for uploads?",
         "Accepted formats depend on the document rule set by the admin for each document. The upload form "
         "will clearly reject unsupported formats. Common accepted types are: PDF, JPG, PNG, DOCX."),
        ("My document was rejected — what do I do?",
         "Open the order, go to the Documents tab, read the rejection reason left by the employee, "
         "prepare a corrected file, and re-upload it. If the order status is 'Waiting for Customer', "
         "use the Missing Documents Response page to submit the new file."),
        ("I forgot my order number — how do I find it?",
         "Log in to your account and open My Orders. Your order number is shown in the first column "
         "of every order row (format: KH2026-XXXXXX)."),
        ("Can I have multiple orders at the same time?",
         "Yes. There is no limit on the number of simultaneous orders you can have in the system."),
        ("Why is my order stuck on 'Waiting for Customer' for a long time?",
         "This status means you have a missing documents request that has not been answered. "
         "Open the order and go to the Missing Documents section to upload the required files."),
        ("Can I cancel an order?",
         "Contact the reviewing employee or support team. Orders in 'New' or 'Under Review' status "
         "may be cancellable. Orders already assigned to a provider may incur a fee."),
        ("I changed the interface language but content still shows in Arabic — why?",
         "Some content (service descriptions, category names) is bilingual. If the admin has only "
         "entered the Arabic version, the Arabic text will show regardless of language setting. "
         "Contact the admin to add English content."),
        ("How do I change my login email?",
         "Email changes require administrator action. Contact your admin with your request."),
        ("A service I need is not in the catalog — what do I do?",
         "Use the Missing Service Request feature (accessible from the services page or your dashboard) "
         "to submit a request for the service to be added. Admins review these requests regularly."),
        ("Can I download my order documents after it is completed?",
         "Yes. Open the completed order from My Orders and scroll to the Final Documents section. "
         "Click Download. Links are valid for 30 minutes; refresh the page to get a new link if needed."),
        ("As an admin, how do I deactivate a user without deleting them?",
         "Go to Rule Management > Users & Roles, find the user, click Edit, "
         "and toggle 'User is active' to OFF. The user cannot log in but all their data is preserved."),
        ("As an admin, how do I reset another user's password?",
         "Go to Rule Management > Users & Roles, click Edit on the user, and enter a new password "
         "in the Password field. Save the record and share the new temporary password with the user."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        rq = p.add_run(f"Q:  {q}")
        rq.font.bold = True; rq.font.size = Pt(10.5); rq.font.color.rgb = DARK_NAVY
        body(doc, f"A:  {a}")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CH 10 — GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════

def ch10(doc):
    h1(doc, "10.  Glossary")
    table(doc, [
        ("Order",                 "A single service request submitted by a customer and tracked through the system."),
        ("Order Number",          "Unique alphanumeric code (e.g. KH2026-000001) that identifies every order."),
        ("Document Rule",         "Admin-configured specification of which files must accompany a service application."),
        ("Assignment Rule",       "Admin-configured link between a service and a provider for automatic order routing."),
        ("Workflow Rule",         "Automated trigger that fires a configured action when an order reaches a certain state."),
        ("Provider",              "An organisation or individual authorised to execute a specific service type."),
        ("Employee",              "Internal staff member responsible for reviewing orders and verifying documents."),
        ("Support",               "Customer-facing staff who can view orders and manage service catalog entries."),
        ("Admin",                 "Super-user with full configuration access across the entire platform."),
        ("JWT",                   "JSON Web Token — secure credential used to authenticate each user session."),
        ("SessionStorage",        "Browser storage where your login token is kept. Cleared when the tab is closed."),
        ("Category",              "A hierarchical grouping label for services in the public catalog."),
        ("Slug",                  "URL-safe lowercase string derived from a name (e.g. 'civil-status-certificates')."),
        ("Audit Log",             "An immutable, timestamped record of every significant action in the system."),
        ("Notification Template", "A reusable message blueprint sent automatically when a defined event occurs."),
        ("Placeholder",           "{variable} tokens inside notification templates, replaced with real values at send time."),
        ("Service Relation",      "A dependency or recommendation link between two services."),
        ("Missing Document",      "A file flagged by an employee as absent or insufficient on a customer order."),
        ("Final Document",        "A deliverable file uploaded by a provider as the result of completing a service."),
        ("Active Content Version","The homepage content set currently published to all public site visitors."),
        ("Theme",                 "A set of brand colors, logo, and fonts applied to the entire site UI."),
        ("Download Token",        "A 30-minute, single-use signed URL for securely downloading an order document."),
    ], headers=["Term", "Definition"])


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    build_cover(doc)
    build_toc(doc)
    ch1(doc)
    ch2(doc)
    ch3(doc)
    ch4(doc)
    ch5(doc)
    ch6(doc)
    ch7(doc)
    ch8(doc)
    ch9(doc)
    ch10(doc)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
